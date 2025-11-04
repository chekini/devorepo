import os
from docx import Document
from docx.shared import Inches

# Ruta donde están las 6 carpetas
carpeta_principal = r'C:\Users\USER1\MAIN FOLDER\SCREENSHOTS\SS folder'  # ← CAMBIA ESTA RUTA
#r'C:\Users\USER1\MAIN FOLDER\SCREENSHOTS\SS'
# Crea el documento Word
doc = Document()

# Recorre cada subcarpeta dentro de la carpeta principal
for nombre_subcarpeta in sorted(os.listdir(carpeta_principal)):
    ruta_subcarpeta = os.path.join(carpeta_principal, nombre_subcarpeta)

    if os.path.isdir(ruta_subcarpeta):
        # Opcional: agrega un encabezado con el nombre de la subcarpeta
        doc.add_heading(f"Carpeta: {nombre_subcarpeta}", level=1)

        # Recorre los archivos JPG dentro de la subcarpeta
        for archivo in sorted(os.listdir(ruta_subcarpeta)):
            if archivo.lower().endswith('.png'):
                ruta_imagen = os.path.join(ruta_subcarpeta, archivo)
                titulo = os.path.splitext(archivo)[0]

                doc.add_heading(titulo, level=2)
                doc.add_picture(ruta_imagen, width=Inches(5))
                doc.add_paragraph('')

# Guarda el documento
doc.save('imagenes_desde_subcarpetas.docx')
